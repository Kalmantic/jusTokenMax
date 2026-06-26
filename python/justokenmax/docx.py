"""DOCX -> Markdown.

A thin extractor over python-docx. We walk the document body in source order so
headings, paragraphs and tables stay interleaved the way the author wrote them
(python-docx exposes ``doc.paragraphs`` and ``doc.tables`` as separate lists, so
a naive walk would emit all text then all tables and lose the ordering).

The value here is format conversion, not compression: a .docx is an opaque ZIP
the model can't read, and this turns it into searchable, quotable Markdown.
Images, comments, tracked changes and headers/footers are intentionally dropped
— they add complexity and rarely carry the meaning.
"""

from __future__ import annotations

from typing import Tuple

# Safety cap. The Read hook runs this automatically on untrusted files, so we
# bound the disk write the same way pdf.py does.
MAX_OUTPUT_CHARS = 5_000_000

_TRUNCATED = "\n\n> _[justokenmax: output truncated — document exceeds safety caps]_\n"

# Raster images (logos, charts, diagrams) are dropped — we extract text only.
# A .docx has no page structure, so we leave a single document-level marker so a
# reader/agent knows visual content existed and can open the source if it matters,
# rather than losing it silently. Mirrors the per-page marker in pdf.py.
_IMG_NOTE = ("> _[justokenmax: {n} image(s) in this document not extracted — "
             "read the source if visual data matters]_")


def _count_images(doc) -> int:
    """Number of images embedded in the document body (inline or floating).

    Counts image relationships on the main document part — not every image part
    in the package, which would also pick up the docProps thumbnail Word writes
    by default. Fail-open: any access error counts as zero."""
    try:
        return sum(1 for rel in doc.part.rels.values()
                   if rel.reltype.endswith("/image") and not rel.is_external)
    except Exception:
        return 0


def _heading_level(style_name: str) -> int:
    """Markdown heading level for a paragraph style, or 0 for body text."""
    if not style_name:
        return 0
    if style_name == "Title":
        return 1
    if style_name.startswith("Heading "):
        try:
            return min(6, int(style_name.split()[1]))
        except (IndexError, ValueError):
            return 0
    return 0


def _table_to_md(table) -> str:
    """Render a python-docx table as a GitHub Markdown table."""
    rows = [[cell.text.replace("\n", " ").strip() for cell in row.cells]
            for row in table.rows]
    rows = [r for r in rows if any(c for c in r)]
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]
    header, body = rows[0], rows[1:]
    md = ["| " + " | ".join(header) + " |",
          "| " + " | ".join(["---"] * width) + " |"]
    md += ["| " + " | ".join(r) + " |" for r in body]
    return "\n".join(md)


def _iter_blocks(doc):
    """Yield paragraphs and tables in document order.

    python-docx keeps paragraphs and tables in separate collections; the only
    way to recover their interleaved order is to walk the body's XML children.
    """
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def docx_to_markdown(path: str) -> Tuple[str, dict]:
    """Return (markdown, {"paragraphs": int, "images": int}). Paragraphs counts
    non-empty text blocks (headings included, since a heading is a paragraph in a
    .docx); images counts embedded image parts, which are dropped but flagged."""
    from docx import Document

    doc = Document(path)
    parts = []
    paragraphs = 0
    images = _count_images(doc)
    chars = 0
    truncated = False

    for block in _iter_blocks(doc):
        if chars > MAX_OUTPUT_CHARS:
            truncated = True
            break
        if hasattr(block, "rows"):  # a Table (paragraphs have no .rows)
            md = _table_to_md(block)
            if md:
                chars += len(md)
                parts.append(md)
                parts.append("")
        else:  # a Paragraph
            text = block.text.strip()
            if not text:
                continue  # skip empty paragraphs
            paragraphs += 1
            level = _heading_level(block.style.name if block.style else "")
            line = ("#" * level + " " + text) if level else text
            chars += len(line)
            parts.append(line)
            parts.append("")

    if images:
        parts.append(_IMG_NOTE.format(n=images))

    out = "\n".join(parts).strip() + "\n"
    if len(out) > MAX_OUTPUT_CHARS:
        out = out[:MAX_OUTPUT_CHARS] + _TRUNCATED
        truncated = True
    elif truncated:
        out = out.rstrip("\n") + _TRUNCATED
    return out, {"paragraphs": paragraphs, "images": images}
