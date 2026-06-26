"""DOCX -> Markdown extraction (optimize kind="docx").

Builds a real .docx with python-docx (no binary fixture committed, same policy
as the PDF/PNG fixtures in conftest) and asserts the Markdown contract.
"""

import json

import pytest

# python-docx is the optional `office` extra; skip this module (don't error) when
# it isn't installed, matching the tiktoken skip convention in test_tokens.py.
pytest.importorskip("docx")

from justokenmax.docx import docx_to_markdown
from justokenmax.optimize import optimize


def _build_docx(path, with_table=True):
    from docx import Document

    doc = Document()
    doc.add_heading("Quarterly Report", level=1)
    doc.add_paragraph("Revenue grew across all regions.")
    doc.add_paragraph("")  # empty paragraph — must be skipped
    doc.add_heading("Regional Breakdown", level=2)
    doc.add_paragraph("APAC led the quarter.")
    if with_table:
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Region"
        table.cell(0, 1).text = "Revenue"
        table.cell(1, 0).text = "APAC"
        table.cell(1, 1).text = "1200"
    doc.add_paragraph("Closing remarks follow the table.")
    doc.save(str(path))


def test_headings_map_to_markdown(tmp_path):
    p = tmp_path / "report.docx"
    _build_docx(p)
    md, stats = docx_to_markdown(str(p))
    assert "# Quarterly Report" in md
    assert "## Regional Breakdown" in md


def test_body_text_and_paragraph_count(tmp_path):
    p = tmp_path / "report.docx"
    _build_docx(p)
    md, stats = docx_to_markdown(str(p))
    assert "Revenue grew across all regions." in md
    assert "APAC led the quarter." in md
    # 4 non-empty paragraphs (the empty one is skipped, headings not counted as body paras)
    assert stats["paragraphs"] >= 4


def test_empty_paragraphs_skipped(tmp_path):
    p = tmp_path / "report.docx"
    _build_docx(p)
    md, _ = docx_to_markdown(str(p))
    # no run of 3+ blank lines
    assert "\n\n\n" not in md


def test_table_rendered_as_markdown(tmp_path):
    p = tmp_path / "report.docx"
    _build_docx(p)
    md, _ = docx_to_markdown(str(p))
    assert "| Region | Revenue |" in md
    assert "| --- | --- |" in md
    assert "| APAC | 1200 |" in md


def test_document_order_preserved(tmp_path):
    """Heading, then table, then the post-table paragraph — in source order.

    python-docx exposes doc.paragraphs and doc.tables as SEPARATE lists, so a
    naive walk would emit all paragraphs then all tables, losing interleaving.
    """
    p = tmp_path / "report.docx"
    _build_docx(p)
    md, _ = docx_to_markdown(str(p))
    i_heading = md.index("Regional Breakdown")
    i_table = md.index("| Region | Revenue |")
    i_after = md.index("Closing remarks follow the table.")
    assert i_heading < i_table < i_after


# ---- optimize() dispatch integration ----

def test_optimize_detects_and_converts_docx(tmp_path):
    p = tmp_path / "report.docx"
    _build_docx(p)
    res = optimize(str(p))
    assert res.ok and res.kind == "docx"
    assert res.output.endswith(".docx.md")
    # Honest accounting: conversion-not-compression, so no token saving claimed.
    assert res.tokens_before == res.tokens_after


def test_secret_in_docx_is_masked(tmp_path):
    from docx import Document

    secret = "AK" + "IA" + "S" * 16  # AWS access key id shape, assembled at runtime
    doc = Document()
    doc.add_paragraph("Deployment notes")
    doc.add_paragraph(f"aws_key = {secret}")
    p = tmp_path / "leaky.docx"
    doc.save(str(p))

    res = optimize(str(p))
    assert res.ok
    artifact = open(res.output, encoding="utf-8").read()
    assert secret not in artifact  # _redact must have masked it


def test_fail_open_when_library_missing(tmp_path, monkeypatch):
    """A missing python-docx must skip (fail-open), never crash the Read hook."""
    from justokenmax import docx as docx_mod

    def _raise_import(_path):
        raise ImportError("No module named 'docx'")

    monkeypatch.setattr(docx_mod, "docx_to_markdown", _raise_import)
    p = tmp_path / "report.docx"
    _build_docx(p)
    res = optimize(str(p))
    assert res.ok is False and res.kind == "skip"
    assert "python-docx" in res.note


def test_fail_open_on_parse_error(tmp_path, monkeypatch):
    """A corrupt/unparseable .docx skips rather than raising."""
    from justokenmax import docx as docx_mod

    def _raise(_path):
        raise ValueError("bad zip")

    monkeypatch.setattr(docx_mod, "docx_to_markdown", _raise)
    p = tmp_path / "report.docx"
    _build_docx(p)
    res = optimize(str(p))
    assert res.ok is False and res.kind == "skip"


def test_cli_docx_subcommand(tmp_path, capsys):
    from justokenmax.cli import main

    p = tmp_path / "report.docx"
    _build_docx(p)
    rc = main(["docx", "--json", str(p)])
    out = json.loads(capsys.readouterr().out)
    assert rc == 0
    assert out["ok"] is True and out["kind"] == "docx"


# ---- dropped-image marker (mirrors the PDF handler / PR #35) ----

def test_flags_images_not_extracted(tmp_path):
    # Raster images are dropped from the text extract; the marker tells the
    # reader visual content existed so it isn't lost silently.
    from docx import Document
    from PIL import Image

    img_path = tmp_path / "pic.png"
    Image.new("RGB", (12, 12), (200, 30, 30)).save(str(img_path))
    doc = Document()
    doc.add_paragraph("Architecture overview")
    doc.add_picture(str(img_path))
    p = tmp_path / "withimage.docx"
    doc.save(str(p))

    md, stats = docx_to_markdown(str(p))
    assert "image(s) in this document not extracted" in md
    assert stats["images"] >= 1


def test_no_image_marker_for_text_only_docx(tmp_path):
    p = tmp_path / "report.docx"
    _build_docx(p)  # no images
    md, stats = docx_to_markdown(str(p))
    assert "not extracted" not in md
    assert stats["images"] == 0
